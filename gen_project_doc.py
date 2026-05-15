"""
Generate a comprehensive DOCX explaining the Student Performance Prediction project.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_DIR = r"d:\Student-performance-prediction-main\Student-performance-prediction-main"
OUTPUT_PATH = os.path.join(PROJECT_DIR, "Student_Performance_Prediction_Project_Documentation.docx")

doc = Document()

# ── STYLES ─────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)

for s in ["Heading 1", "Heading 2", "Heading 3"]:
    hs = doc.styles[s]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(55, 48, 163)

# ── HELPER FUNCTIONS ───────────────────────────
def add_code_block(doc, code_text, caption=""):
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_section_break(doc):
    doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════
doc.add_paragraph("")
doc.add_paragraph("")

title = doc.add_heading("Student Performance Prediction System", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Complete Project Documentation")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph("")
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run(
    "A comprehensive, point-by-point guide explaining every file, "
    "function, and line of code in simple words with real-world examples."
)
run.font.size = Pt(12)
run.font.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Project Overview & Purpose",
    "2. Technologies Used",
    "3. Project Folder Structure",
    "4. File-by-File Explanation",
    "   4.1 train_model.py (Machine Learning Engine)",
    "   4.2 app.py (Flask Web Server & Backend Logic)",
    "   4.3 templates/upload.html (Frontend UI & Chatbot)",
    "   4.4 templates/base.html (Master Layout Template)",
    "   4.5 Other Supporting Files",
    "5. How the System Works (Step-by-Step Data Flow)",
    "6. Key Machine Learning Concepts Explained",
    "7. Summary",
]
for item in toc_items:
    if item.startswith("   "):
        add_bullet(doc, item.strip())
    else:
        doc.add_paragraph(item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("1. Project Overview & Purpose", level=1)

doc.add_paragraph(
    "This project is an AI-powered web application called StudentLens that predicts "
    "whether a student will PASS or FAIL their exams. It uses Machine Learning algorithms "
    "trained on historical student data (study habits, absences, family support, etc.) "
    "to make these predictions."
)

doc.add_heading("Why is this useful?", level=3)
doc.add_paragraph(
    "Imagine a teacher with 300 students. It is impossible to manually track every "
    "student's risk factors. This system automates that work. It instantly identifies "
    "at-risk students and suggests specific interventions like 'increase study time' "
    "or 'reduce absences'."
)

doc.add_heading("Real-World Example", level=3)
t = doc.add_table(rows=3, cols=4)
t.style = "Light Grid Accent 1"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
hdr[0].text = "Student"
hdr[1].text = "Study Time"
hdr[2].text = "Absences"
hdr[3].text = "AI Prediction"
row1 = t.rows[1].cells
row1[0].text = "Student A"
row1[1].text = "4 hrs/day"
row1[2].text = "0"
row1[3].text = "PASS (Low Risk)"
row2 = t.rows[2].cells
row2[0].text = "Student B"
row2[1].text = "1 hr/day"
row2[2].text = "8"
row2[3].text = "FAIL (High Risk)"

add_section_break(doc)

# ══════════════════════════════════════════════════════════════════════
# 2. TECHNOLOGIES USED
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("2. Technologies Used", level=1)

tech_items = [
    ("Python 3.x", " - The main programming language for both the ML model and the web server."),
    ("Flask", " - A lightweight Python web framework that serves HTML pages and handles HTTP requests."),
    ("scikit-learn", " - A machine learning library providing algorithms like Logistic Regression, SVM, and Random Forest."),
    ("XGBoost", " - An advanced gradient boosting algorithm, often the best performer on tabular data."),
    ("Pandas & NumPy", " - Libraries for data manipulation (reading CSVs, computing statistics)."),
    ("Matplotlib", " - Used for generating charts (confusion matrices, ROC curves, bar charts)."),
    ("FPDF", " - A library for generating PDF reports programmatically."),
    ("Bootstrap 5", " - A CSS framework for styling the web UI with responsive, modern design."),
    ("JavaScript", " - Handles the interactive chatbot, dynamic tabs, and AJAX form submission in the browser."),
    ("HTML/CSS", " - The fundamental web technologies for structure and styling."),
]
for name, desc_text in tech_items:
    add_bullet(doc, desc_text, bold_prefix=name)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 3. PROJECT FOLDER STRUCTURE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("3. Project Folder Structure", level=1)
doc.add_paragraph("Here is what each file and folder in the project does:")

folder_code = """Student-performance-prediction-main/
|
|-- app.py                  # The Flask web server (backend)
|-- train_model.py          # The ML training script
|-- model.pkl               # The saved trained AI model (binary file)
|-- algo_metrics.json       # Saved accuracy/metrics of all algorithms
|
|-- uploads/
|   |-- student_data.csv    # The historical student dataset
|
|-- outputs/
|   |-- prediction_*.csv    # Generated prediction result files
|
|-- static/
|   |-- charts/             # Auto-generated PNG chart images
|       |-- roc_curve.png
|       |-- confusion_matrix_raw.png
|       |-- algo_comparison.png
|       |-- feature_importance.png
|       |-- ...
|
|-- templates/
|   |-- base.html           # Master HTML layout (navbar, fonts, styles)
|   |-- upload.html          # Main page (chatbot, upload, analytics tabs)
|
|-- requirements.txt        # Python package dependencies
"""
add_code_block(doc, folder_code, "Folder Tree:")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 4. FILE-BY-FILE EXPLANATION
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("4. File-by-File Explanation", level=1)

# ── 4.1 train_model.py ──
doc.add_heading("4.1 train_model.py  -  The Machine Learning Engine", level=2)
doc.add_paragraph(
    "This file is the 'brain builder'. You run it ONCE to teach the AI from historical data. "
    "After running, it saves the trained model as model.pkl. Here is a section-by-section breakdown:"
)

doc.add_heading("Section 1: Imports (Lines 1-34)", level=3)
doc.add_paragraph(
    "The script starts by importing all necessary libraries. Think of this as "
    "'gathering your tools before starting work'."
)
add_code_block(doc, """import numpy as np               # For mathematical operations (arrays, means)
import pandas as pd              # For reading CSV files into tables
from sklearn.model_selection import train_test_split  # Splits data into train/test
from sklearn.ensemble import RandomForestClassifier   # One of the ML algorithms
from xgboost import XGBClassifier                     # Another powerful ML algorithm
import joblib                    # For saving the trained model to a file
import matplotlib.pyplot as plt  # For drawing charts""", "Key Imports Explained:")

doc.add_paragraph(
    "Example: numpy is like a calculator, pandas is like Excel inside Python, "
    "and sklearn is the 'machine learning toolbox'."
)

doc.add_heading("Section 2: Loading the Dataset (Lines 36-52)", level=3)
add_code_block(doc, """data = pd.read_csv("uploads/student_data.csv")
y = data["passed"].map({"no": 0, "yes": 1})   # Convert "yes"/"no" to 1/0
X = data.drop(columns=["passed"])              # Everything EXCEPT the answer""", "Code:")

doc.add_paragraph(
    "In simple words: We read the CSV file. The column 'passed' is the ANSWER we want to predict "
    "(Pass=1, Fail=0). Everything else (study time, absences, etc.) are the QUESTIONS (features) "
    "we give to the AI."
)
doc.add_paragraph(
    "Example: If the CSV has 395 rows, that means 395 students. Each row has about 30 columns "
    "like age, study time, family support, etc."
)

doc.add_heading("Section 3: Feature Types & Preprocessing (Lines 54-73)", level=3)
doc.add_paragraph("The data has two types of columns:")
add_bullet(doc, ' "yes"/"no", "urban"/"rural" - these are text (categorical). The AI cannot understand text, '
           'so we convert them to numbers using OneHotEncoder.', bold_prefix="Categorical Features:")
add_bullet(doc, " age=17, absences=5 - these are already numbers, so we pass them through directly.", bold_prefix="Numeric Features:")

add_code_block(doc, """preprocess = ColumnTransformer(
    transformers=[
        ("cat", OneHotEncoder(handle_unknown="ignore"), categorical_features),
        ("num", "passthrough", numeric_features),
    ]
)""", "Code:")

doc.add_paragraph(
    'Example: If a student\'s "school" is "GP", OneHotEncoder converts it to [1, 0]. '
    'If it is "MS", it becomes [0, 1]. This way, the AI sees numbers, not letters.'
)

doc.add_heading("Section 4: Handling Class Imbalance (Lines 75-83)", level=3)
doc.add_paragraph(
    "If 300 students passed and only 80 failed, the AI might just always guess 'Pass' and get 79% accuracy "
    "while being completely wrong about failing students. To fix this, we calculate 'class weights' "
    "that tell the AI: 'Pay MORE attention to failing students because they are rare but important.'"
)
add_code_block(doc, """class_weights = compute_class_weight("balanced", classes=[0, 1], y=y)
# Result example: {0: 1.8, 1: 0.6}
# This means: A failing student counts as 1.8x more important than a passing one.""", "Code:")

doc.add_heading("Section 5: Defining 5 Algorithms (Lines 85-116)", level=3)
doc.add_paragraph("The system trains 5 different 'AI students' and picks the smartest one:")

alg_table = doc.add_table(rows=6, cols=3)
alg_table.style = "Light Grid Accent 1"
alg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
alg_hdr = alg_table.rows[0].cells
alg_hdr[0].text = "Algorithm"
alg_hdr[1].text = "Type"
alg_hdr[2].text = "Simple Explanation"
alg_data = [
    ("Logistic Regression", "Linear", "Draws a straight line to separate Pass/Fail students."),
    ("SVM (Linear)", "Linear", "Finds the BEST straight line with the widest gap between groups."),
    ("SVM (RBF)", "Non-Linear", "Draws a curved boundary for complex patterns."),
    ("Random Forest", "Ensemble", "Creates 200 decision trees and takes a vote from all of them."),
    ("XGBoost", "Boosting", "Trains 300 small trees sequentially, each fixing the mistakes of the last one."),
]
for i, (name, typ, explanation) in enumerate(alg_data, start=1):
    row = alg_table.rows[i].cells
    row[0].text = name
    row[1].text = typ
    row[2].text = explanation

doc.add_heading("Section 6: Train/Test Split (Lines 118-121)", level=3)
add_code_block(doc, """X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)""", "Code:")
doc.add_paragraph(
    "We split the data: 80% for training (teaching the AI) and 20% for testing (checking if it learned well). "
    "The 'stratify=y' ensures both groups have the same pass/fail ratio."
)
doc.add_paragraph(
    "Example: If we have 395 students, 316 go to training and 79 go to testing."
)

doc.add_heading("Section 7: Cross-Validation & Training Loop (Lines 123-217)", level=3)
doc.add_paragraph("For each of the 5 algorithms, the code:")
add_bullet(doc, "Performs 5-fold cross-validation (splits data 5 different ways and averages the scores).")
add_bullet(doc, "Trains the model on the training set.")
add_bullet(doc, "Tests it on the test set and records Accuracy, Precision, Recall, and F1-Score.")
add_bullet(doc, "Keeps track of which algorithm scored best.")

doc.add_paragraph(
    "At the end, the best algorithm (e.g., 'XGBoost with 70.13% CV accuracy') is saved."
)

doc.add_heading("Section 8: Saving the Model & Generating Charts (Lines 226-334)", level=3)
add_code_block(doc, 'joblib.dump(best_pipeline, "model.pkl")', "Saving the model:")
doc.add_paragraph("The trained model is saved as model.pkl. This file contains the entire 'brain' of the AI.")
doc.add_paragraph("The script then generates several charts:")
add_bullet(doc, "Confusion Matrix (Counts) - Shows correct vs incorrect predictions.")
add_bullet(doc, "Confusion Matrix (Normalized) - Shows percentages of correct predictions per class.")
add_bullet(doc, "Algorithm Comparison Bar Chart - Compares all 5 algorithms side by side.")
add_bullet(doc, "Feature Importance - Shows which factors matter most (e.g., absences, study time).")
add_bullet(doc, "ROC Curve (All Models) - Compares model performance with True Positive vs False Positive rates.")

doc.add_page_break()

# ── 4.2 app.py ──
doc.add_heading("4.2 app.py  -  The Flask Web Server (Backend)", level=2)
doc.add_paragraph(
    "This is the largest file (~1564 lines). It is the 'traffic controller' that connects "
    "the website to the AI model. Every time a user clicks a button, app.py decides what happens."
)

doc.add_heading("Section 1: Imports & Setup (Lines 1-58)", level=3)
add_code_block(doc, """from flask import Flask, render_template, request, send_file
import pandas as pd
import joblib
from fpdf import FPDF

app = Flask(__name__)
MODEL = joblib.load("model.pkl")   # Load the trained AI brain""", "Key Code:")
doc.add_paragraph(
    "In simple words: We import the tools, create a Flask web server, and load the saved AI model. "
    "Any time someone visits the website, this 'app' object handles the request."
)

doc.add_heading("Section 2: create_charts() Function (Lines 61-276)", level=3)
doc.add_paragraph(
    "This function takes a DataFrame of student predictions and generates 9 different charts. "
    "Each chart is saved as a .png image in the static/charts/ folder."
)
chart_items = [
    ("Pass/Fail Count Bar Chart", "Shows how many students passed vs failed."),
    ("Study Time by Prediction", "Average study time for pass vs fail groups."),
    ("Absences Histogram", "Distribution of how many students have how many absences."),
    ("Failures by Prediction", "Average past failures for pass vs fail."),
    ("Correlation Heatmap", "Shows which features are mathematically related to each other."),
    ("Confusion Matrix (Counts)", "Loaded from the training phase."),
    ("Confusion Matrix (Normalized)", "Percentage-based confusion matrix."),
    ("Algorithm Comparison", "Bar chart comparing all 5 algorithm accuracies."),
    ("ROC Curve (All Models)", "Trade-off between true positive and false positive rates for all models."),
]
ct = doc.add_table(rows=len(chart_items)+1, cols=2)
ct.style = "Light Grid Accent 1"
ct.rows[0].cells[0].text = "Chart Name"
ct.rows[0].cells[1].text = "What It Shows"
for i, (cn, cd) in enumerate(chart_items, start=1):
    ct.rows[i].cells[0].text = cn
    ct.rows[i].cells[1].text = cd

doc.add_heading("Section 3: generate_feedback() Function (Lines 279-407)", level=3)
doc.add_paragraph(
    "This function generates human-readable feedback by analyzing patterns in the data. "
    "For example, if students predicted to fail have an average of 7.2 absences but those predicted to "
    "pass have only 2.1, it writes: 'Higher absence counts are associated with a higher risk of failing.'"
)
doc.add_paragraph(
    "It also identifies 'Swing Students' - those with a 40-60% probability - who are on the border "
    "and most likely to benefit from intervention."
)

doc.add_heading("Section 4: Flask Routes (Lines 564-856)", level=3)
doc.add_paragraph("Routes are the URLs the user visits. Each route triggers a specific action:")

route_table = doc.add_table(rows=5, cols=3)
route_table.style = "Light Grid Accent 1"
rh = route_table.rows[0].cells
rh[0].text = "URL"
rh[1].text = "Method"
rh[2].text = "What It Does"
routes = [
    ("/", "GET", "Loads the main page with default predictions."),
    ("/", "POST", "Accepts a new CSV upload, runs predictions, shows results."),
    ("/predict_manual", "POST", "Chatbot sends student data here; returns pass/fail prediction."),
    ("/download", "GET", "Downloads the CSV file with all predictions."),
]
for i, (url, method, desc_r) in enumerate(routes, start=1):
    route_table.rows[i].cells[0].text = url
    route_table.rows[i].cells[1].text = method
    route_table.rows[i].cells[2].text = desc_r

doc.add_heading("Section 5: get_individual_insights() Function (Lines 679-781)", level=3)
doc.add_paragraph(
    "This is the 'smart advisor' function. For a single student, it analyzes each factor and "
    "calculates a numeric 'impact' score from -100 (terrible) to +100 (excellent)."
)
add_code_block(doc, """factors = [
    {"id": "studytime", "label": "Study Effort", "inverse": False, ...},
    {"id": "absences",  "label": "Attendance",    "inverse": True, ...},
    {"id": "failures",  "label": "Academic Foundation", "inverse": True, ...},
    ...
]""", "Simplified Logic:")
doc.add_paragraph(
    "For each factor, if 'inverse' is True (like absences), a HIGH value means BAD. "
    "If inverse is False (like study time), a HIGH value means GOOD."
)
doc.add_paragraph(
    "The function then sorts factors by impact and classifies them as Strengths (+40 and above) "
    "or Risks (below 0). For each risk, it generates an actionable tip."
)
doc.add_paragraph(
    "Example output for a student with 8 absences:\n"
    "  - Risk: 'Attendance' (impact = -60)\n"
    "  - Tip: 'Maintain an attendance rate above 95%.'"
)

doc.add_heading("Section 6: PDF Report Generation (Lines 887-1564)", level=3)
doc.add_paragraph(
    "When a user clicks 'Download Report', the app creates a professional PDF using the FPDF library. "
    "The custom PredictionPDF class adds a branded header (indigo bar with 'StudentLens' title) "
    "and footer (page number + date) to every page."
)
doc.add_paragraph("The PDF includes:")
add_bullet(doc, "A cover page with the project title and generation date.")
add_bullet(doc, "A 'Key Input Values' table listing the student's data.")
add_bullet(doc, "A 'Factor Impact Analysis' section with visual bars showing positive/negative impacts.")
add_bullet(doc, "A 'Personalized Action Plan' section with specific tips and interventions.")
add_bullet(doc, "An 'Analytical Visualizations' appendix embedding all generated charts.")

doc.add_page_break()

# ── 4.3 upload.html ──
doc.add_heading("4.3 templates/upload.html  -  The Frontend (User Interface)", level=2)
doc.add_paragraph(
    "This is the file the user actually SEES in their browser. It contains HTML structure, "
    "CSS styling, and JavaScript interactivity. It uses Bootstrap 5 for the dark theme."
)

doc.add_heading("Part 1: Three Main Tabs", level=3)
tab_table = doc.add_table(rows=4, cols=2)
tab_table.style = "Light Grid Accent 1"
tab_table.rows[0].cells[0].text = "Tab Name"
tab_table.rows[0].cells[1].text = "Purpose"
tabs = [
    ("Upload Dataset", "Upload a CSV file of many students to get batch predictions."),
    ("Interactive Assessment (Chatbot)", "Answer questions one-by-one for a single student prediction."),
    ("Analytics", "View all generated charts and model comparison information."),
]
for i, (tn, tp) in enumerate(tabs, start=1):
    tab_table.rows[i].cells[0].text = tn
    tab_table.rows[i].cells[1].text = tp

doc.add_heading("Part 2: The Chatbot (JavaScript)", level=3)
doc.add_paragraph("The chatbot is implemented entirely in JavaScript within upload.html. It works as follows:")
add_bullet(doc, "An array of 11 questions is defined (study time, failures, absences, family support, etc.).")
add_bullet(doc, "Questions are displayed one at a time as animated chat bubbles.")
add_bullet(doc, "The user answers by clicking option buttons or typing a number.")
add_bullet(doc, "Each answer is saved into a hidden HTML form field.")
add_bullet(doc, "After all 11 questions, a summary card appears showing all answers.")
add_bullet(doc, "When the user clicks 'Submit', the hidden form is sent to /predict_manual via HTTP POST.")

doc.add_heading("Part 3: Results Display", level=3)
doc.add_paragraph("When the prediction comes back from the server, JavaScript dynamically:")
add_bullet(doc, "Shows the prediction ('PASS' or 'FAIL') with a colored badge.")
add_bullet(doc, "Shows the confidence percentage (e.g., '79%').")
add_bullet(doc, "Displays a Risk Level badge: GREEN (Low Risk), AMBER (Medium Risk), RED (High Risk).")
add_bullet(doc, "Lists Key Strengths (green + icons) and Critical Risks (red - icons) in circular badges.")
add_bullet(doc, "Draws horizontal 'impact bars' for each factor showing positive or negative influence.")

doc.add_heading("Part 4: Styling", level=3)
doc.add_paragraph("The page uses a dark, glass-morphism aesthetic with:")
add_bullet(doc, "Dark background gradients (deep blue to black).")
add_bullet(doc, "Semi-transparent glass cards with blur effects (backdrop-filter: blur).")
add_bullet(doc, "Smooth hover animations on buttons and cards.")
add_bullet(doc, "Color-coded elements: green for strengths, red for risks, indigo for headers.")

doc.add_page_break()

# ── 4.4 base.html ──
doc.add_heading("4.4 templates/base.html  -  Master Layout Template", level=2)
doc.add_paragraph("This file contains the common HTML structure shared by all pages. It includes:")
add_bullet(doc, "The <head> section with CSS links (Bootstrap, Google Fonts, custom styles).")
add_bullet(doc, "The navigation bar at the top.")
add_bullet(doc, "A {% block content %} placeholder where upload.html injects its content.")
add_bullet(doc, "Footer scripts (Bootstrap JS, custom JavaScript).")
doc.add_paragraph(
    "This 'template inheritance' pattern avoids duplicating the same navbar/footer code on every page."
)

# ── 4.5 Other Files ──
doc.add_heading("4.5 Other Supporting Files", level=2)
other_files = [
    ("model.pkl", " - The saved trained model. Generated by train_model.py. Loaded by app.py."),
    ("algo_metrics.json", " - A JSON file containing accuracy, precision, recall, and F1-score for all 5 algorithms."),
    ("uploads/student_data.csv", " - The historical dataset with ~395 student records and 30+ feature columns."),
    ("requirements.txt", " - Lists all Python packages needed: flask, pandas, scikit-learn, xgboost, fpdf2, matplotlib, etc."),
    ("static/charts/", " - Folder containing all auto-generated PNG chart images."),
]
for fname, fdesc in other_files:
    add_bullet(doc, fdesc, bold_prefix=fname)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 5. STEP-BY-STEP DATA FLOW
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("5. How the System Works (Step-by-Step Data Flow)", level=1)
doc.add_heading("Scenario: User predicts a single student via Chatbot", level=2)

steps = [
    ("Step 1: User Opens Website",
     "The browser loads upload.html (via base.html). The 'Interactive Assessment' tab appears with the chatbot."),
    ("Step 2: User Answers 11 Questions",
     "The JavaScript chatbot asks questions like 'How many hours does the student study daily?' "
     "The user clicks buttons (e.g., '3-4 hours'). Each answer fills a hidden form field."),
    ("Step 3: Form Submission",
     "When the user clicks 'Submit', all answers are packaged as an HTTP POST request "
     "and sent to the /predict_manual URL on the Flask server."),
    ("Step 4: Server Receives Data",
     "The predict_manual() function in app.py catches the data. It converts each field to the correct "
     "type (numbers stay numbers, text stays text) and builds a Pandas DataFrame."),
    ("Step 5: Model Prediction",
     "app.py loads model.pkl and calls MODEL.predict_proba(). The model outputs [0.21, 0.79], "
     "meaning 21% chance of failing and 79% chance of passing."),
    ("Step 6: Insight Generation",
     "get_individual_insights() analyzes each factor. For example, absences=8 produces a negative "
     "impact score of -60, generating a 'Risk' label and the tip 'Maintain attendance above 95%'."),
    ("Step 7: Response Sent to Browser",
     "Flask renders upload.html with the prediction data injected. The browser receives the HTML."),
    ("Step 8: Results Displayed",
     "JavaScript reads the injected data and dynamically shows the Pass/Fail badge, confidence score, "
     "risk level, strengths list, risks list, and impact bars."),
    ("Step 9: PDF Download (Optional)",
     "If the user clicks 'Download Report', a request hits /download_manual_pdf. "
     "PredictionPDF creates a PDF with tables, impact bars, and tips, and sends it for download."),
]

for title_s, desc_s in steps:
    doc.add_heading(title_s, level=3)
    doc.add_paragraph(desc_s)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 6. KEY ML CONCEPTS
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("6. Key Machine Learning Concepts Explained", level=1)

concepts = [
    ("Accuracy",
     "The percentage of correct predictions out of all predictions. "
     "Example: If the model correctly classifies 70 out of 100 students, accuracy is 70%."),
    ("Precision",
     "Of all students the model predicts as 'Pass', how many actually passed? "
     "Example: Model says 80 will pass, but only 60 do. Precision = 60/80 = 75%."),
    ("Recall",
     "Of all students who actually passed, how many did the model correctly identify? "
     "Example: 70 actually passed, model found 60. Recall = 60/70 = 85.7%."),
    ("F1-Score",
     "The harmonic mean of Precision and Recall. It balances both metrics into one number. "
     "High F1 = both precision and recall are good."),
    ("Confusion Matrix",
     "A 2x2 grid showing: True Positives (correctly predicted Pass), True Negatives "
     "(correctly predicted Fail), False Positives (predicted Pass but Fail), "
     "and False Negatives (predicted Fail but Pass)."),
    ("ROC Curve",
     "A graph plotting True Positive Rate vs False Positive Rate at various thresholds. "
     "A curve closer to the top-left corner means better performance. AUC near 1.0 = excellent, 0.5 = random."),
    ("Cross-Validation",
     "Instead of testing once, the data is split 5 different ways and tested 5 times. "
     "The average score is more reliable than a single test."),
    ("Class Imbalance",
     "When one category has far fewer examples than the other. "
     "Without correction, the model might always guess the majority class. Class weights fix this."),
    ("Feature Importance",
     "A ranking of which input variables matter most for the prediction. "
     "For example, 'absences' might be the #1 most important feature."),
    ("OneHotEncoding",
     "Converting text values to numbers. "
     "Example: 'school' with values GP/MS becomes two columns: school_GP and school_MS (0 or 1)."),
]

for cname, cdesc in concepts:
    doc.add_heading(cname, level=3)
    doc.add_paragraph(cdesc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 7. SUMMARY
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("7. Summary", level=1)
doc.add_paragraph(
    "This project is a complete, end-to-end AI system for student performance prediction. "
    "It covers every stage from data loading, to model training, to web deployment, to PDF reporting."
)

summary_items = [
    "train_model.py LEARNS from historical data by training 5 algorithms and saves the best one.",
    "app.py SERVES the website, receives user data, runs predictions, and generates insights.",
    "upload.html DISPLAYS the interactive chatbot, results, and analytics to the user.",
    "base.html provides the common layout skeleton.",
    "The PDF generator creates professional, downloadable reports.",
    "The ROC Curve, Confusion Matrices, and charts provide transparency into how the AI works.",
]
for item in summary_items:
    add_bullet(doc, item)

doc.add_paragraph("")
final = doc.add_paragraph()
run = final.add_run(
    "With this system, educators can identify at-risk students early and take targeted actions "
    "to improve academic outcomes - powered by the intelligence of machine learning."
)
run.italic = True

# ── SAVE ──
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
